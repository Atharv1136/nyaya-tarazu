"""
backend/routers/export.py
==========================
POST /export — generate a Word (.docx) document from a structured brief.
"""

import io
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from backend.models.schemas import ExportRequest, SingleBrief

router = APIRouter(prefix="/export", tags=["export"])


def _brief_to_docx(brief: SingleBrief, case_title: Optional[str] = None) -> bytes:
    """Convert a SingleBrief to a .docx file bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()

    # ---- Style setup ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # ---- Title block ----
    title = doc.add_heading("", level=0)
    title_run = title.add_run(
        case_title or f"Legal Brief — {brief.persona.capitalize()} Side"
    )
    title_run.font.size = Pt(16)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Persona and date
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared for: {brief.persona.upper()} COUNSEL\n").bold = True
    meta.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")

    doc.add_paragraph()  # spacer

    # ---- Disclaimer ----
    disclaimer_para = doc.add_paragraph()
    disclaimer_run = disclaimer_para.add_run(f"⚠ {brief.disclaimer}")
    disclaimer_run.italic = True
    disclaimer_run.font.size = Pt(10)
    disclaimer_run.font.color.rgb = RGBColor(0x7A, 0x2E, 0x2E)  # oxblood

    doc.add_paragraph()

    # ---- Issues ----
    if brief.issues:
        doc.add_heading("Issues", level=1)
        for i, issue in enumerate(brief.issues, 1):
            doc.add_paragraph(f"{i}. {issue}", style="List Number")

    # ---- Applicable Provisions ----
    if brief.applicable_provisions:
        doc.add_heading("Applicable Provisions", level=1)
        for section in brief.applicable_provisions:
            h = doc.add_heading(section.heading, level=2)
            doc.add_paragraph(section.content)
            if section.citations:
                cit_para = doc.add_paragraph()
                cit_run = cit_para.add_run("Citations: " + ", ".join(section.citations))
                cit_run.font.name = "Courier New"
                cit_run.font.size = Pt(10)

    # ---- Arguments ----
    if brief.arguments:
        doc.add_heading("Arguments", level=1)
        for section in brief.arguments:
            doc.add_heading(section.heading, level=2)
            doc.add_paragraph(section.content)
            if section.citations:
                cit_para = doc.add_paragraph()
                cit_run = cit_para.add_run("Citations: " + ", ".join(section.citations))
                cit_run.font.name = "Courier New"
                cit_run.font.size = Pt(10)

    # ---- Precedents ----
    if brief.supporting_precedents:
        doc.add_heading("Supporting Precedents", level=1)
        for section in brief.supporting_precedents:
            doc.add_heading(section.heading, level=2)
            doc.add_paragraph(section.content)

    # ---- Prayer ----
    if brief.prayer:
        doc.add_heading("Prayer / Relief Sought", level=1)
        doc.add_paragraph(brief.prayer)

    # ---- Unverified citations ----
    unverified = [c for c in (brief.citation_verifications or []) if not c.verified]
    if unverified:
        doc.add_heading("Citations Requiring Manual Verification", level=1)
        for c in unverified:
            p = doc.add_paragraph()
            r = p.add_run(f"• {c.citation}: {c.flag_reason or 'Verify manually.'}")
            r.font.color.rgb = RGBColor(0x7A, 0x2E, 0x2E)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.post("")
async def export_endpoint(request: ExportRequest) -> StreamingResponse:
    """
    Export a legal brief as a Word document (.docx).
    Returns a file download response.
    """
    if request.format != "docx":
        raise HTTPException(status_code=422, detail="Only 'docx' format is currently supported.")

    try:
        docx_bytes = _brief_to_docx(request.brief, request.case_title)
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(exc)[:300]}")

    persona = request.brief.persona
    filename = f"nyaya_tarazu_{persona}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
